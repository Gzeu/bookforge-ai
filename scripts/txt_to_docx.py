#!/usr/bin/env python3
"""BookForge AI — TXT to DOCX export for paperback workflows."""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt

CHAPTER_RE = re.compile(r"(Chapter\s+\d+[:\s—–\-]?[^\n]*|CHAPTER\s+\d+[^\n]*)", re.IGNORECASE)


def detect_chapters(text: str):
    parts = CHAPTER_RE.split(text)
    if len(parts) > 1:
        chapters = []
        for i in range(1, len(parts), 2):
            title = parts[i].strip()
            body = parts[i + 1].strip() if i + 1 < len(parts) else ""
            chapters.append((title, body))
        return chapters
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunk = max(1, len(paragraphs) // 10)
    return [(f"Chapter {idx+1}", "\n".join(paragraphs[i:i+chunk])) for idx, i in enumerate(range(0, len(paragraphs), chunk))]


def txt_to_docx(input_file: str, output_file: str, title: str = "My Novel", author: str = "Author Name") -> str:
    text = Path(input_file).read_text(encoding="utf-8")
    chapters = detect_chapters(text)
    doc = Document()
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    doc.add_paragraph(author)
    doc.add_page_break()

    for chapter_title, chapter_body in chapters:
        h = doc.add_paragraph()
        r = h.add_run(chapter_title)
        r.bold = True
        r.font.size = Pt(18)
        for line in [ln.strip() for ln in chapter_body.split("\n") if ln.strip()]:
            if line in ("* * *", "***", "---"):
                p = doc.add_paragraph("* * *")
                p.alignment = 1
            else:
                doc.add_paragraph(line)
        doc.add_page_break()

    Path(output_file).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_file)
    return output_file
